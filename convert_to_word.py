"""
Convert HACKATHON_JOURNEY.md to Word document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title page
    title = doc.add_heading('Bob ICA Catalysts Hackathon Journey', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Travel Refund Uncertainty Estimation System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Team: Abhijit Joshi\n').bold = True
    meta.add_run('Event: Bob ICA Catalysts May 2026 Hackathon\n').bold = True
    meta.add_run('Date: May 21-23, 2026\n').bold = True
    meta.add_run('Project: AI-Powered Travel Refund Estimation with IBM ICA Integration').bold = True
    
    doc.add_page_break()
    
    # Read markdown content
    with open('HACKATHON_JOURNEY.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse and add content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip title page content
        if i < 10 and ('Bob ICA Catalysts Hackathon Journey' in line or 
                       'Travel Refund Uncertainty Estimation System' in line or
                       line.startswith('**Team:**') or 
                       line.startswith('**Event:**') or 
                       line.startswith('**Date:**') or 
                       line.startswith('**Project:**')):
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 128)
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove markdown formatting
            text = text.replace('**', '')
            text = text.replace('`', '')
            p = doc.add_paragraph(text, style='List Bullet')
            if '✅' in text:
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif '❌' in text:
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        elif line[0].isdigit() and line[1:3] in ['. ', ') ']:
            text = line[3:]
            text = text.replace('**', '')
            text = text.replace('`', '')
            p = doc.add_paragraph(text, style='List Number')
        
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Bold paragraphs
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        
        # Regular paragraphs
        elif line.strip():
            # Remove markdown formatting
            text = line.replace('**', '')
            text = text.replace('`', '')
            
            p = doc.add_paragraph(text)
            
            # Color code checkmarks
            if '✅' in text:
                for run in p.runs:
                    if '✅' in run.text:
                        run.font.color.rgb = RGBColor(0, 128, 0)
            elif '❌' in text:
                for run in p.runs:
                    if '❌' in run.text:
                        run.font.color.rgb = RGBColor(255, 0, 0)
        
        i += 1
    
    # Save document
    doc.save('HACKATHON_JOURNEY.docx')
    print('[SUCCESS] Word document created successfully: HACKATHON_JOURNEY.docx')
    print('[INFO] Document contains complete hackathon journey with all 8 steps')
    print('[INFO] Location: c:/Users/AbhijitJoshi/Uncertainty_Refund/HACKATHON_JOURNEY.docx')

if __name__ == '__main__':
    main()

# Made with Bob
