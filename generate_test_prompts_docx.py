"""
Generate Word document with test prompts for hackathon demo
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_test_prompts_document():
    """Create a professional Word document with test prompts"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Travel Refund Estimation Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Test Prompts & Scenarios for Hackathon Demo', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run('Agent Name: ').bold = True
    meta.add_run('TravelRefundAdvisor\n')
    meta.add_run('Purpose: ').bold = True
    meta.add_run('AI-powered travel refund estimation during force majeure events\n')
    meta.add_run('Date: ').bold = True
    meta.add_run(f'{datetime.now().strftime("%B %d, %Y")}\n')
    meta.add_run('Event: ').bold = True
    meta.add_run('Bob ICA Catalysts May 2026 Hackathon')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Introduction
    doc.add_heading('Overview', 1)
    intro = doc.add_paragraph(
        'This document contains 10 comprehensive test scenarios to demonstrate the '
        'Travel Refund Estimation Assistant during your hackathon presentation. '
        'Each scenario is designed to showcase different capabilities of the AI agent '
        'and the underlying ML-powered refund estimation system.'
    )
    
    doc.add_page_break()
    
    # Scenario 1
    doc.add_heading('Scenario 1: Basic Booking Query', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('Show me all my travel bookings')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Agent retrieves list of bookings', style='List Bullet')
    doc.add_paragraph('• Displays booking references, destinations, dates', style='List Bullet')
    doc.add_paragraph('• Shows total costs', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Shows basic data retrieval capability')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph()  # Spacing
    
    # Scenario 2
    doc.add_heading('Scenario 2: Specific Booking Details', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('Can you show me the details for booking ID 1?')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Retrieves specific booking information', style='List Bullet')
    doc.add_paragraph('• Shows customer details', style='List Bullet')
    doc.add_paragraph('• Lists all components (flights, hotels, visas, insurance)', style='List Bullet')
    doc.add_paragraph('• Displays individual costs', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Demonstrates detailed data access')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_page_break()
    
    # Scenario 3
    doc.add_heading('Scenario 3: Refund Estimation - High Severity', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph(
        'I need to cancel my booking to Paris (booking ID 1) due to a natural disaster. '
        'Can you estimate my refund with high severity?'
    )
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Calls estimate_refund API', style='List Bullet')
    doc.add_paragraph('• Shows expected refund amount and percentage', style='List Bullet')
    doc.add_paragraph('• Displays 95% confidence interval (lower and upper bounds)', style='List Bullet')
    doc.add_paragraph('• Provides best case, worst case, and most likely scenarios', style='List Bullet')
    doc.add_paragraph('• Shows risk score', style='List Bullet')
    doc.add_paragraph('• Explains force majeure considerations', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('⭐ Core ML-powered refund estimation feature')
    demo_value.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    demo_value.runs[0].bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Scenario 4
    doc.add_heading('Scenario 4: Global Risk Assessment', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('What are the current global risk events affecting travel?')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Lists active force majeure events worldwide', style='List Bullet')
    doc.add_paragraph('• Shows event types (war, pandemic, natural disaster, etc.)', style='List Bullet')
    doc.add_paragraph('• Displays severity levels', style='List Bullet')
    doc.add_paragraph('• Indicates affected regions', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Real-time risk monitoring capability')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_page_break()
    
    # Scenario 5
    doc.add_heading('Scenario 5: Regional Risk Check', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph("I'm planning to travel to Ukraine. What's the current risk level there?")
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Retrieves regional risk assessment', style='List Bullet')
    doc.add_paragraph('• Shows aggregate risk level (low/medium/high/critical)', style='List Bullet')
    doc.add_paragraph('• Lists specific events affecting the region', style='List Bullet')
    doc.add_paragraph('• Provides travel safety recommendations', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Location-specific risk analysis')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph()  # Spacing
    
    # Scenario 6
    doc.add_heading('Scenario 6: Refund Statistics Analysis', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('Show me historical refund statistics by component type and event type')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Displays aggregated refund statistics', style='List Bullet')
    doc.add_paragraph('• Shows average refund percentages by component and event type', style='List Bullet')
    doc.add_paragraph('• Indicates force majeure success rates', style='List Bullet')
    doc.add_paragraph('• Provides insights on refund patterns', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Data-driven insights and trends')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_page_break()
    
    # Scenario 7
    doc.add_heading('Scenario 7: Provider Policy Inquiry', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('What are the refund policies for different travel providers?')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Lists provider refund policies', style='List Bullet')
    doc.add_paragraph('• Shows standard vs. force majeure refund percentages', style='List Bullet')
    doc.add_paragraph('• Displays cancellation fees', style='List Bullet')
    doc.add_paragraph('• Compares different providers', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Policy comparison and transparency')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph()  # Spacing
    
    # Scenario 8
    doc.add_heading('Scenario 8: Complex Multi-Component Booking', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph(
        'I have a booking with flights, hotel, and visa. If I cancel due to political unrest '
        'with medium severity, what would be my expected refund for each component?'
    )
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Analyzes each component separately', style='List Bullet')
    doc.add_paragraph('• Provides component-specific refund estimates', style='List Bullet')
    doc.add_paragraph('• Shows total expected refund', style='List Bullet')
    doc.add_paragraph('• Explains differences in refund rates by component type', style='List Bullet')
    doc.add_paragraph('• Considers provider-specific policies', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Sophisticated multi-component analysis')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_page_break()
    
    # Scenario 9
    doc.add_heading('Scenario 9: Confidence Interval Explanation', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph('Can you explain what the confidence intervals mean in my refund estimate?')
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Explains 95% confidence interval concept', style='List Bullet')
    doc.add_paragraph('• Clarifies lower and upper bounds', style='List Bullet')
    doc.add_paragraph('• Describes uncertainty in predictions', style='List Bullet')
    doc.add_paragraph('• Relates to historical data patterns', style='List Bullet')
    doc.add_paragraph('• Helps customer understand risk', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('⭐ AI transparency and explainability')
    demo_value.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    demo_value.runs[0].bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Scenario 10
    doc.add_heading('Scenario 10: Comparative Scenario Analysis', 1)
    doc.add_heading('Prompt:', 2)
    p = doc.add_paragraph(
        'Compare the refund estimates for booking ID 1 under low, medium, and high severity scenarios'
    )
    p.style = 'Intense Quote'
    
    doc.add_heading('Expected Response:', 2)
    doc.add_paragraph('• Generates multiple estimates with different severity levels', style='List Bullet')
    doc.add_paragraph('• Shows side-by-side comparison', style='List Bullet')
    doc.add_paragraph('• Highlights differences in expected refunds', style='List Bullet')
    doc.add_paragraph('• Explains impact of severity on refund amounts', style='List Bullet')
    doc.add_paragraph('• Helps customer make informed decisions', style='List Bullet')
    
    doc.add_heading('Demo Value:', 2)
    demo_value = doc.add_paragraph('Decision support and scenario planning')
    demo_value.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_page_break()
    
    # Quick Test Prompts
    doc.add_heading('Additional Quick Test Prompts', 1)
    doc.add_paragraph('Use these for quick functionality checks:', style='List Bullet')
    doc.add_paragraph()
    
    quick_prompts = [
        'Are you connected to the refund estimation system?',
        'What can you help me with?',
        "What's the refund policy for Air India?",
        'Show me historical refund data for flight cancellations during pandemics',
        'Show me only active risk events'
    ]
    
    for i, prompt in enumerate(quick_prompts, 1):
        p = doc.add_paragraph(f'{i}. "{prompt}"', style='List Number')
    
    doc.add_paragraph()
    
    # Demo Flow
    doc.add_heading('Recommended 5-Minute Demo Flow', 1)
    
    doc.add_heading('Minute 1: Introduction (30 seconds)', 2)
    doc.add_paragraph('• Introduce the AI-powered travel refund estimation assistant', style='List Bullet')
    doc.add_paragraph('• Mention ML models for predicting refunds during force majeure', style='List Bullet')
    
    doc.add_heading('Minute 2: Basic Functionality (1 minute)', 2)
    doc.add_paragraph('• Test: "Show me all my travel bookings"', style='List Bullet')
    doc.add_paragraph('• Test: "Show me details for booking ID 1"', style='List Bullet')
    
    doc.add_heading('Minute 3: Core Feature (1.5 minutes)', 2)
    doc.add_paragraph('• Test: "Estimate refund for booking 1 with high severity"', style='List Bullet')
    doc.add_paragraph('• Highlight: Confidence intervals, scenarios, ML predictions', style='List Bullet')
    
    doc.add_heading('Minute 4: Risk Assessment (1 minute)', 2)
    doc.add_paragraph('• Test: "What are current global risk events?"', style='List Bullet')
    doc.add_paragraph('• Test: "What\'s the risk level for Ukraine?"', style='List Bullet')
    
    doc.add_heading('Minute 5: Advanced Features (1 minute)', 2)
    doc.add_paragraph('• Test: "Show me refund statistics"', style='List Bullet')
    doc.add_paragraph('• Test: "Compare low vs high severity scenarios"', style='List Bullet')
    doc.add_paragraph('• Wrap up with business value', style='List Bullet')
    
    doc.add_page_break()
    
    # Key Points
    doc.add_heading('Key Points to Emphasize', 1)
    
    doc.add_heading('Technical Excellence:', 2)
    tech_points = [
        'Full-stack application (React + FastAPI)',
        'ML-powered predictions (Random Forest, Gradient Boosting)',
        'Real-time API integration',
        'IBM ICA agent integration',
        'Confidence intervals and uncertainty quantification'
    ]
    for point in tech_points:
        doc.add_paragraph(f'✅ {point}', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Business Value:', 2)
    business_points = [
        'Reduces customer service workload',
        'Provides instant refund estimates',
        'Improves customer satisfaction',
        'Data-driven decision making',
        'Transparent AI explanations'
    ]
    for point in business_points:
        doc.add_paragraph(f'✅ {point}', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Innovation:', 2)
    innovation_points = [
        'AI agent for travel industry',
        'Force majeure event handling',
        'Uncertainty estimation (not just point predictions)',
        'Multi-component booking analysis',
        'Real-time risk monitoring'
    ]
    for point in innovation_points:
        doc.add_paragraph(f'✅ {point}', style='List Bullet')
    
    doc.add_page_break()
    
    # Troubleshooting
    doc.add_heading('Troubleshooting Tips', 1)
    
    doc.add_heading('If Agent Doesn\'t Respond:', 2)
    doc.add_paragraph('1. Check backend is running: python main.py', style='List Number')
    doc.add_paragraph('2. Check ngrok is active', style='List Number')
    doc.add_paragraph('3. Verify API health endpoint', style='List Number')
    
    doc.add_heading('If Data Seems Wrong:', 2)
    doc.add_paragraph('1. Database has synthetic data for demo', style='List Number')
    doc.add_paragraph('2. Refund estimates are ML predictions (not actual)', style='List Number')
    doc.add_paragraph('3. Risk events are sample data', style='List Number')
    
    doc.add_heading('If Agent Says "Can\'t Access":', 2)
    doc.add_paragraph('1. Ensure OpenAPI tools are configured', style='List Number')
    doc.add_paragraph('2. Check ngrok URL hasn\'t changed', style='List Number')
    doc.add_paragraph('3. Verify API endpoints are accessible', style='List Number')
    
    doc.add_paragraph()
    
    # Final Checklist
    doc.add_heading('Final Checklist Before Demo', 1)
    checklist = [
        'Backend running (python main.py)',
        'Ngrok active and URL noted',
        'Agent responding in ICA',
        'Tested at least 3 prompts',
        'Prepared to explain ML models',
        'Ready to discuss business value',
        'Backup slides/screenshots ready',
        'Confident and enthusiastic!'
    ]
    for item in checklist:
        doc.add_paragraph(f'☐ {item}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run('Good luck with your hackathon! 🚀').bold = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer2 = doc.add_paragraph()
    footer2.add_run('Your agent is working perfectly - just follow these prompts and you\'ll have an impressive demo!')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    filename = 'HACKATHON_TEST_PROMPTS.docx'
    doc.save(filename)
    print(f'[SUCCESS] Document created successfully: {filename}')
    print(f'[INFO] Total pages: ~15-20 pages')
    print(f'[INFO] Contains: 10 detailed scenarios + quick tests + demo flow + troubleshooting')
    
    return filename

if __name__ == '__main__':
    import sys
    # Set UTF-8 encoding for Windows console
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    
    create_test_prompts_document()

# Made with Bob
